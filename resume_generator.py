"""
resume_generator.py

Opens "Darrow Jason FINAL resume.docx" as a formatting template and replaces
only the dynamic sections (summary + Voya bullets) with content from the
tailored JSON produced by resume_tailor.py.  All margins, fonts, colors,
borders, and static sections are inherited from the template unchanged.

Usage:
    python3 resume_generator.py --input output/tailored_Acme.json
    python3 resume_generator.py --input output/tailored_Acme.json --output personal/Acme_Resume.docx
"""

import argparse
import copy
import json
import math
import os

from docx import Document
from docx.oxml.ns import qn

import band

# ── Page-length estimate ───────────────────────────────────────────────────
# A tailored resume should land on two pages.  Nothing in the pipeline enforces
# that: resume_tailor.py puts no length budget on the summary or the bullets,
# so a verbose model response silently produces a three-page resume.
#
# Rather than add a heavyweight dependency (converting to PDF just to count
# pages means every user installs LibreOffice), we estimate from wrapped line
# counts.  The constants below were calibrated against rendered output from
# this template — predicted vs. actual line counts agree within one line:
#
#     tailored_MIT_AICR.json  → predicted 79, actual 78  (2 pages)
#     tailored_Citizens.json  → predicted 86, actual 85  (3 pages)
#     tailored_Lumen.json     → predicted 90, actual 90  (3 pages)
#
# This is a warning, not a gate.  The .docx is always written — you edit and
# review it in Word before sending.  The warning just makes an overlong resume
# impossible to miss.

CHARS_PER_LINE_BODY   = 105   # full text width
CHARS_PER_LINE_BULLET = 110   # indented, but a smaller font offsets the indent
LINES_PER_PAGE        = 40
TARGET_PAGES          = 2


def estimate_line_count(document) -> int:
    """Estimate how many rendered lines the document occupies."""
    lines = 0
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue
        try:
            is_bullet = para.style.name == "List Paragraph"
        except Exception:  # noqa: BLE001 — style can be None on odd documents
            is_bullet = False
        width = CHARS_PER_LINE_BULLET if is_bullet else CHARS_PER_LINE_BODY
        lines += max(1, math.ceil(len(text) / width))
    return lines


def report_page_estimate(document) -> None:
    """Print an estimated page count, loudly if it runs past the target."""
    lines = estimate_line_count(document)
    budget = LINES_PER_PAGE * TARGET_PAGES
    pages = max(1, math.ceil(lines / LINES_PER_PAGE))

    if lines > budget:
        over = lines - budget
        print("", flush=True)
        print("=" * 68, flush=True)
        print(f"⚠  LENGTH WARNING — this resume is likely {pages} pages, not {TARGET_PAGES}.", flush=True)
        print(f"   Roughly {over} line(s) over a {TARGET_PAGES}-page layout.", flush=True)
        print("   Open it in Word and trim before sending.  Usual culprits:", flush=True)
        print("     • a long tailored_summary (aim for ~650 characters)", flush=True)
        print("     • long bullets (aim for ~230 characters each)", flush=True)
        print("     • more Voya bullets than the intended 5", flush=True)
        print("=" * 68, flush=True)
    elif lines > budget - 6:
        print(
            f"• Length: ~{pages} page(s), close to the {TARGET_PAGES}-page limit "
            f"({lines}/{budget} lines) — check the last page before sending.",
            flush=True,
        )
    else:
        print(f"• Length: ~{pages} page(s) ({lines}/{budget} lines).", flush=True)

# ── Args ───────────────────────────────────────────────────────────────────

parser = argparse.ArgumentParser(description="Generate tailored resume Word doc")
parser.add_argument(
    "--input",
    required=True,
    help="Tailored JSON file (e.g. output/tailored_Acme.json)",
)
parser.add_argument(
    "--output",
    default=os.path.join("personal", "Jason_Darrow_Resume.docx"),
    help="Output filename (default: personal/Jason_Darrow_Resume.docx)",
)
parser.add_argument(
    "--template",
    default=os.path.join("personal", "Darrow Jason FINAL resume.docx"),
    help="Template .docx to use as the formatting base",
)
parser.add_argument(
    "--early-experience",
    type=int,
    default=0,
    metavar="N",
    help="Print N bullets from the early-career role (default 0, off).  Use "
         "only when a posting genuinely needs what that role proves.",
)
parser.add_argument(
    "--early-role",
    default="Clicktolearn.com",
    help="Company name or experience id for --early-experience",
)
args = parser.parse_args()

os.makedirs(os.path.dirname(args.output) or ".", exist_ok=True)
print(f"Template : {args.template}", flush=True)
print(f"Input    : {args.input}", flush=True)

# ── Load data ──────────────────────────────────────────────────────────────

with open(args.input) as f:
    tailored = json.load(f)

# ── Open template ──────────────────────────────────────────────────────────

doc = Document(args.template)

# ── Helpers ────────────────────────────────────────────────────────────────

def replace_para_text(para, new_text):
    """
    Replace all text in a paragraph's runs with new_text, preserving the
    formatting (font, size, color, bold, italic) of the first run.
    """
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def remove_para(para):
    """Remove a paragraph element from the document."""
    para._element.getparent().remove(para._element)


def set_xml_para_text(p_elem, text):
    """
    Set the visible text of a raw lxml <w:p> element.
    Clears all runs then sets the first run's <w:t> to text.
    """
    runs = p_elem.findall(".//" + qn("w:r"))
    if not runs:
        return
    for r in runs:
        for t in r.findall(qn("w:t")):
            t.text = ""
    ts = runs[0].findall(qn("w:t"))
    if ts:
        ts[0].text = text
        # xml:space="preserve" prevents Word from stripping leading/trailing spaces
        ts[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


# ── 0. Strip template-specific page-break layout from generated docs ───────
# The FINAL resume has a hardcoded page break inside the BoA role description
# paragraph that forces a "Page Two" header at a fixed position.  That only
# works when content fills page 1 exactly.  In generated resumes the content
# length varies, so we remove the break and the "Page Two" header line and
# let Word paginate naturally.

PAGE_BREAK_ANCHOR  = "Served as Application Manager"
PAGE_TWO_HEADER    = "Page Two"
BOA_CONTINUED      = "BANK OF AMERICA (Continued)"

# Walk a copy of the list — we may remove items mid-iteration
for para in list(doc.paragraphs):
    txt = para.text.strip()

    # Remove the run containing the page break from the BoA role description
    if PAGE_BREAK_ANCHOR in txt:
        for r in para._element.findall(".//" + qn("w:r")):
            for br in r.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    r.getparent().remove(r)
                    print("✓ Hardcoded page break removed", flush=True)
                    break

    # Remove the "Page Two" header paragraph entirely
    if PAGE_TWO_HEADER in txt:
        remove_para(para)
        print("✓ Page Two header removed", flush=True)

    # Remove the "BANK OF AMERICA (Continued)" label — redundant without page break
    if txt == BOA_CONTINUED:
        remove_para(para)
        print("✓ BoA (Continued) label removed", flush=True)

# ── 1. Replace the summary paragraph ──────────────────────────────────────
# The summary is the first paragraph whose text starts with the known anchor.

SUMMARY_ANCHOR = "IT Delivery Manager with experience"

for para in doc.paragraphs:
    if para.text.strip().startswith(SUMMARY_ANCHOR):
        replace_para_text(para, tailored["tailored_summary"])
        print("✓ Summary replaced", flush=True)
        break
else:
    print("⚠ Summary paragraph not found — check SUMMARY_ANCHOR", flush=True)

# ── 1a. Rewrite the headline and tagline ──────────────────────────────────
# "IT Leader" and the three terms under it were static in identity.title, so
# every resume this pipeline ever produced opened with the same positioning --
# Dell, MIT, Nitto, Intact, SSi and Attensi all said "Program & Project
# Delivery | Vendor & Portfolio Management | IT Governance".
#
# Worse, on the Attensi resume all three duplicated the Core Competencies band
# fourteen words below: Program & Project Delivery verbatim, IT Governance
# minus a word, Portfolio Management inside Vendor & Portfolio Management.
# Three lines of the best real estate on the page carrying no new information.
#
# One deliberate overlap is fine and is emphasis -- the core identity term
# earns a second mention.  Three is an echo, so the generator says so rather
# than silently editing, since which term to repeat is a judgement call.

HEADLINE_DEFAULT = "IT Leader"
HEADLINE_SEP     = " | "


def replace_headline(document, tailored_json, master_file):
    head_cfg = (master_file or {}).get("headline") or {}
    title    = tailored_json.get("tailored_headline") or head_cfg.get("title")
    tagline  = tailored_json.get("tailored_tagline") or head_cfg.get("tagline") or []
    if not title and not tagline:
        return

    template_title   = head_cfg.get("title") or HEADLINE_DEFAULT
    template_tagline = HEADLINE_SEP.join(head_cfg.get("tagline") or [])

    title_para = tagline_para = None
    for para in document.paragraphs:
        txt = para.text.strip()
        if txt == template_title and title_para is None:
            title_para = para
        elif title_para is not None and txt == template_tagline:
            tagline_para = para
            break

    if title_para is None:
        print("⚠ Headline paragraph not found — template headline kept", flush=True)
        return

    if title and title != template_title:
        replace_para_text(title_para, title)
    if tagline and tagline_para is not None:
        replace_para_text(tagline_para, HEADLINE_SEP.join(tagline))
    elif tagline:
        print("⚠ Tagline paragraph not found — headline updated, tagline kept", flush=True)

    # Echo check against the competency band.
    comps = tailored_json.get("tailored_competencies") or []
    echoed = [t for t, _ in
              [(t, None) for t in tagline]
              if band.dedupe_skills([t], comps)[1]]
    if len(echoed) >= len(tagline) and tagline:
        print(
            f"⚠ Tagline: all {len(tagline)} terms repeat the Core Competencies band "
            f"({', '.join(echoed)}).  Swap at least one for something the band "
            f"does not already say.",
            flush=True,
        )
    else:
        print(
            f"✓ Headline: {title or template_title} — "
            + HEADLINE_SEP.join(tagline or head_cfg.get("tagline", []))
            + (f"  ({len(echoed)} deliberate repeat)" if echoed else ""),
            flush=True,
        )


try:
    with open("master_resume.json") as _fh:
        master_data = json.load(_fh)
except (OSError, ValueError):
    master_data = {}

replace_headline(doc, tailored, master_data)

# ── 1b. Rewrite the Core Competencies band ────────────────────────────────
# This band was static template text on every resume ever generated -- Dell,
# MIT, Nitto and Intact all shipped the same ten phrases.  resume_tailor.py has
# produced a per-role key_skills list since Wave 2 and nothing ever wrote it to
# the document, so the densest keyword real estate on the page said the same
# thing to every employer.
#
# It is also where the skills terms now go.  The template has no skills
# section at all, which meant a posting asking for cloud, GitHub or DevOps had
# nowhere on the page those words could legitimately appear -- the reason the
# Intact resume could not pass a requirements screen no matter which bullets
# were selected.
#
# Three defects showed up on the Attensi resume and are fixed here:
#
#   1. The skills row was prefixed "Technical: ".  It read as a label the rest
#      of the band does not have, and half the terms on it were not technical
#      (Project Management, Financial Services).
#   2. The skills row repeated the competencies directly above it -- Stakeholder
#      Management appeared twice, Project/Program Management restated Program &
#      Project Delivery, Agile and Waterfall restated Agile & Waterfall
#      Delivery.  Repetition in the densest keyword block on the page reads as
#      padding to a human and adds nothing for a scanner.
#   3. The band inherits the Normal style, which is JUSTIFIED.  A row long
#      enough to wrap gets its word spacing stretched edge to edge, which is
#      what "the spacing is off" looked like.  Every band row is now written
#      left-aligned with consistent space-before, and long rows are split at a
#      character budget so they never wrap in the first place.

COMPETENCY_HEADING = "Core Competencies"
COMPETENCY_END     = "Professional Experience"
BAND_SPACE_BEFORE  = 60   # twips (3pt), matching the template's band rows


def band_lines(tailored_json):
    """Build the replacement text lines for the competency band."""
    comps = tailored_json.get("tailored_competencies") or []
    per_line = band.COMPETENCIES_PER_LINE
    lines = [
        " | ".join(comps[i:i + per_line])
        for i in range(0, len(comps), per_line)
    ]

    skills, dropped = band.dedupe_skills(tailored_json.get("technical_skills") or [], comps)
    if dropped:
        print(
            f"• Skills row: dropped {len(dropped)} term(s) already in the "
            f"competencies above ({', '.join(dropped)})",
            flush=True,
        )
    lines += band.wrap_terms(skills)
    return [ln for ln in lines if ln.strip()]


def format_band_para(para_or_elem):
    """
    Force left alignment and consistent space-before on a band row.

    The band has no pPr of its own in the template, so it inherits Normal's
    JUSTIFY.  Left-aligning it means a row that does wrap degrades into a
    normal ragged second line instead of a stretched one.
    """
    elem = getattr(para_or_elem, "_element", para_or_elem)
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = elem.makeelement(qn("w:pPr"), {})
        elem.insert(0, pPr)

    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = pPr.makeelement(qn("w:jc"), {})
        pPr.append(jc)
    jc.set(qn("w:val"), "left")

    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = pPr.makeelement(qn("w:spacing"), {})
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(BAND_SPACE_BEFORE))


new_lines = band_lines(tailored)

if not new_lines:
    print("• Core Competencies: no tailored competencies in JSON — template text kept", flush=True)
else:
    in_band = False
    band_paras = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt == COMPETENCY_HEADING:
            in_band = True
            continue
        if in_band:
            if txt == COMPETENCY_END:
                break
            if txt:
                band_paras.append(para)

    if not band_paras:
        print("⚠ Core Competencies band not found — check COMPETENCY_HEADING", flush=True)
    else:
        n_existing, n_new = len(band_paras), len(new_lines)
        for i, para in enumerate(band_paras):
            if i < n_new:
                replace_para_text(para, new_lines[i])
                format_band_para(para)
            else:
                remove_para(para)

        if n_new > n_existing:
            # Clone the last band paragraph for the overflow lines so the new
            # rows inherit the band's font and colour.
            template_elem = band_paras[-1]._element
            last_elem = template_elem
            for text in new_lines[n_existing:]:
                new_p = copy.deepcopy(template_elem)
                set_xml_para_text(new_p, text)
                format_band_para(new_p)
                last_elem.addnext(new_p)
                last_elem = new_p

        n_comps = len(tailored.get("tailored_competencies") or [])
        n_skills = len(tailored.get("technical_skills") or [])
        print(
            f"✓ Core Competencies rewritten: {n_comps} competencies"
            + (f" + {n_skills} skill term(s)" if n_skills else " (no skills row)")
            + f", {n_new} row(s)",
            flush=True,
        )

# ── 2. Replace Voya Financial bullets ─────────────────────────────────────
# Collect the tailored Voya bullets from the JSON.

selected_voya = [
    b["bullet"]
    for b in tailored.get("selected_bullets", [])
    if b.get("company") == "Voya Financial"
]

# Walk paragraphs to find:
#   - all existing List Paragraph bullets under the "IT Delivery Manager" role
#   - the BANK OF AMERICA paragraph that follows them (insertion anchor)

in_voya_role = False
voya_bullet_paras = []
boa_para = None

for para in doc.paragraphs:
    txt = para.text.strip()
    style = para.style.name

    # The role line is a Normal paragraph with exactly "IT Delivery Manager"
    if style == "Normal" and txt == "IT Delivery Manager":
        in_voya_role = True
        continue

    if in_voya_role:
        if style == "Normal" and txt.startswith("BANK OF AMERICA"):
            boa_para = para
            in_voya_role = False
            break
        if style == "List Paragraph" and txt:
            voya_bullet_paras.append(para)

if not voya_bullet_paras:
    print("⚠ Voya bullet paragraphs not found", flush=True)
elif not boa_para:
    print("⚠ BANK OF AMERICA anchor paragraph not found", flush=True)
else:
    n_existing = len(voya_bullet_paras)
    n_selected = len(selected_voya)

    if n_selected == 0:
        print("⚠ No Voya bullets in tailored JSON — template bullets left unchanged", flush=True)
    elif n_selected <= n_existing:
        # Replace the first n_selected paragraphs; remove the rest
        for i, para in enumerate(voya_bullet_paras):
            if i < n_selected:
                replace_para_text(para, selected_voya[i])
            else:
                remove_para(para)
        print(
            f"✓ Voya bullets: replaced {n_selected}, removed {n_existing - n_selected}",
            flush=True,
        )
    else:
        # Fill all existing paragraphs, then clone extras before boa_para
        for i, para in enumerate(voya_bullet_paras):
            replace_para_text(para, selected_voya[i])

        template_bullet_elem = voya_bullet_paras[0]._element
        last_elem = voya_bullet_paras[-1]._element  # track insertion point
        for bullet_text in selected_voya[n_existing:]:
            new_p = copy.deepcopy(template_bullet_elem)
            set_xml_para_text(new_p, bullet_text)
            # addnext inserts immediately after the last bullet, before any spacer
            last_elem.addnext(new_p)
            last_elem = new_p  # advance insertion point

        print(
            f"✓ Voya bullets: replaced {n_existing}, added {n_selected - n_existing}",
            flush=True,
        )

# ── 3. Replace Bank of America bullets ────────────────────────────────────
# resume_tailor.py has selected 2-3 BoA bullets on every run since Wave 2 and
# this script never wrote them.  The BoA section came from the template every
# time, so the selection was decorative -- and worse, check_coverage() scored
# those unwritten bullets as if they were on the page.  The Attensi run is the
# clean example: the tailor picked "...facilitating requirements and test
# planning sessions across multiple lines of business", which answers that
# posting's workshops-and-facilitation must-have, and it never printed.
#
# The BoA section holds three roles, so a bullet has to land under the role it
# actually belongs to.  A 2012-2018 bullet printed under the 2001-2008 heading
# is a factual error, not a formatting one.  Each bullet is resolved back to
# its role through master_resume.json; the JSON's own title field is the
# fallback, and the model does sometimes invent a merged title there
# ("Service Delivery Consultant / Tech Team Manager"), which is why the master
# file is checked first.
#
# A role keeps its template bullets when nothing was selected for it.  Blanking
# a role because the tailor happened to ignore it would lose the resume
# writer's content for no gain.

BOA_COMPANY   = "BANK OF AMERICA"
BOA_SECTION_END = "Education"


def normalize_text(text):
    return " ".join(str(text).lower().split())


def master_bullet_roles():
    """Map normalized master bullet text -> the role title it belongs to."""
    roles = {}
    try:
        with open("master_resume.json") as fh:
            master = json.load(fh)
    except (OSError, ValueError):
        print("• BoA bullets: master_resume.json unreadable — using JSON titles only", flush=True)
        return roles
    for exp in master.get("experience", []):
        for bullet in exp.get("bullets", []):
            roles[normalize_text(bullet.get("text", ""))] = exp.get("title", "")
    return roles


def role_blocks(document, company_prefix, end_heading):
    """
    Split a company's section into {role title: [bullet paragraphs]}.

    Role heading lines are Normal paragraphs containing a tab (the template
    right-aligns the date range on a tab stop).  The role description under a
    heading has no tab, which is what separates the two.
    """
    blocks, order = {}, []
    current, in_company = None, False
    for para in document.paragraphs:
        txt = para.text.strip()
        style = para.style.name
        if style == "Normal" and txt.upper().startswith(company_prefix):
            in_company = True
            continue
        if not in_company:
            continue
        if style == "Normal" and txt == end_heading:
            break
        if style == "Normal" and "\t" in para.text:
            current = para.text.split("\t")[0].strip().rstrip(",")
            blocks.setdefault(current, [])
            order.append(current)
        elif style == "List Paragraph" and txt and current:
            blocks[current].append(para)
    return blocks, order


selected_boa = [
    b for b in tailored.get("selected_bullets", [])
    if b.get("company") == "Bank of America"
]

if not selected_boa:
    print("• BoA bullets: none selected — template bullets kept", flush=True)
else:
    blocks, order = role_blocks(doc, BOA_COMPANY, BOA_SECTION_END)
    if not blocks:
        print("⚠ Bank of America role blocks not found — template bullets kept", flush=True)
    else:
        bullet_roles = master_bullet_roles()

        def resolve_role(entry):
            """Which BoA heading does this bullet belong under?"""
            from_master = bullet_roles.get(normalize_text(entry.get("bullet", "")))
            for candidate in (from_master, entry.get("title")):
                if not candidate:
                    continue
                for heading in order:
                    if normalize_text(heading) == normalize_text(candidate):
                        return heading
            return order[0]  # most recent BoA role

        assigned = {}
        for entry in selected_boa:
            assigned.setdefault(resolve_role(entry), []).append(entry["bullet"])

        replaced = kept = 0
        for heading, paras in blocks.items():
            texts = assigned.get(heading)
            if not texts or not paras:
                kept += len(paras)
                continue
            for i, para in enumerate(paras):
                if i < len(texts):
                    replace_para_text(para, texts[i])
                else:
                    remove_para(para)
            if len(texts) > len(paras):
                template_elem = paras[0]._element
                last_elem = paras[-1]._element
                for text in texts[len(paras):]:
                    new_p = copy.deepcopy(template_elem)
                    set_xml_para_text(new_p, text)
                    last_elem.addnext(new_p)
                    last_elem = new_p
            replaced += len(texts)

        print(
            f"✓ BoA bullets: wrote {replaced} tailored, kept {kept} template "
            f"across {len(blocks)} role(s)",
            flush=True,
        )

# ── 4. Optional: early-career experience ──────────────────────────────────
# The Clicktolearn.com entry (2000-2001) was cut from the template on 7/27 as
# an age signal, and off by default is still the right default.  But it is the
# only external client-facing evidence Jason has, and the only eLearning
# evidence, so for a posting from a training-software company that wants
# client-facing delivery it is the most on-target line in the whole file.
#
# That is a per-application judgement call, so it is a flag rather than a
# permanent template change.  The incremental age cost is smaller than it
# looks: the resume already prints BANK OF AMERICA 2001, so a reader doing the
# arithmetic lands in the same place either way.

if args.early_experience:
    try:
        with open("master_resume.json") as fh:
            master_file = json.load(fh)
    except (OSError, ValueError):
        master_file = {}

    early = None
    for exp in master_file.get("experience", []):
        if exp.get("id") == args.early_role or exp.get("company") == args.early_role:
            early = exp
            break

    anchor = None
    for para in doc.paragraphs:
        if para.style.name == "List Paragraph" and para.text.strip():
            anchor = para
        if para.style.name == "Normal" and para.text.strip() == BOA_SECTION_END:
            break

    company_line = next(
        (p for p in doc.paragraphs
         if p.style.name == "Normal" and p.text.strip().upper().startswith(BOA_COMPANY)),
        None,
    )

    if not early:
        print(f"⚠ Early experience '{args.early_role}' not in master_resume.json", flush=True)
    elif not anchor or not company_line:
        print("⚠ Early experience: insertion anchor not found", flush=True)
    else:
        bullets = sorted(
            early.get("bullets", []),
            key=lambda b: b.get("strength", 0),
            reverse=True,
        )[: args.early_experience]

        # Reuse the BANK OF AMERICA company line's run layout so the entry
        # inherits the bold company name, the tab stop and the date run.
        years = " ".join(str(early.get("dates", "")).split())
        for token in ("January", "February", "March", "April", "May", "June",
                      "July", "August", "September", "October", "November",
                      "December"):
            years = years.replace(token + " ", "")

        heading = copy.deepcopy(company_line._element)
        run_texts = [
            early.get("company", "").upper(), ", ", early.get("title", ""),
            "\t", years,
        ]
        runs = heading.findall(".//" + qn("w:r"))
        for i, run in enumerate(runs):
            for t in run.findall(qn("w:t")):
                t.text = run_texts[i] if i < len(run_texts) else ""
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        # The BANK OF AMERICA line it was cloned from opens a section and has
        # no space above it.  This one follows a bullet, so it needs the same
        # breathing room the sub-role headings get.
        pPr = heading.find(qn("w:pPr"))
        if pPr is None:
            pPr = heading.makeelement(qn("w:pPr"), {})
            heading.insert(0, pPr)
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = pPr.makeelement(qn("w:spacing"), {})
            pPr.append(spacing)
        spacing.set(qn("w:before"), "120")

        anchor._element.addnext(heading)
        last_elem = heading
        for bullet in bullets:
            new_p = copy.deepcopy(anchor._element)
            set_xml_para_text(new_p, bullet.get("text", ""))
            last_elem.addnext(new_p)
            last_elem = new_p

        print(
            f"✓ Early experience added: {early.get('company')} "
            f"({years}), {len(bullets)} bullet(s)",
            flush=True,
        )

# ── Save ───────────────────────────────────────────────────────────────────

doc.save(args.output)
print(f"✓ Resume saved: {args.output}", flush=True)

# Estimate length last so the warning is the final thing on screen.
report_page_estimate(doc)
