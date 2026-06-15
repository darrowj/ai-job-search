"""
cover_letter_generator.py

Generates a tailored cover letter using the Claude API.  Reads the full job
description (pasted text) and optionally the tailored resume JSON produced by
resume_tailor.py to ground the letter in the same bullets and key skills
already selected for this role.

Usage (CLI):
    python3 cover_letter_generator.py --company "Fidelity" --title "IT Delivery Manager"
    python3 cover_letter_generator.py --company "Fidelity" --title "IT Delivery Manager" \\
        --description "$(cat jd.txt)"

The --description flag is required when running from the dashboard (subprocess
cannot use interactive stdin).  If omitted, the script prompts for a manual paste.

Output:
    personal/CoverLetter_COMPANY.docx
"""

import anthropic
import argparse
import json
import os
import re
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()

# ── Context files ────────────────────────────────────────────────────────────

def _load(path):
    try:
        with open(path) as f:
            return f.read()
    except FileNotFoundError:
        return ""

ABOUT_ME      = _load("context/AboutMe.md")
CAREER_HISTORY = _load("context/CareerHistory.md")

# ── Voice and anti-AI rules baked into every prompt ─────────────────────────

VOICE_RULES = """
JASON'S WRITING VOICE — apply every rule without exception:

1. Double space after periods.  Two spaces, not one.  Always.
2. No em dashes or hyphens used as pauses or separators.  Use parentheses,
   commas, or rewrite the sentence instead.  Hyphens inside compound words
   (cross-functional, multi-million-dollar) are fine.
3. Short sentences.  One idea per sentence.  If a sentence runs long, break it.
4. BANNED words and phrases — never use these:
   progressive, leveraging, synergy, results-driven, proven track record,
   passionate about, excited to, dynamic, innovative, strategic thinker,
   seeking to utilize, spearheading, impactful, looking to bring my expertise,
   seasoned professional, extensive background, I am writing to express,
   I would be a great fit, I am the ideal candidate, thrilled, delighted.
5. Give the reason, not just the claim.  Specifics carry the weight.  Numbers
   and context make the case.  "I built an aircraft residual calculator
   supporting $1.2B in business" beats "I have strong technical skills."
6. Direct and confident, not boastful.  State facts.  Let outcomes speak.
7. First person (I, my) is correct for a cover letter.

ANTI-AI RULES — the letter must not feel AI-generated:
- NEVER open with "I am excited to apply", "I am writing to express my interest",
  or any variant.  Start somewhere more interesting.
- NEVER use "I look forward to hearing from you" as the closing line.
  Make the closing specific to the company or the conversation.
- NEVER use filler transitions: Furthermore, Additionally, Moreover,
  In conclusion, In summary, To that end.
- Every sentence must carry a specific fact, number, or concrete action.
  No sentence may exist only to connect two other sentences.
- Do not editorialize.  Never say something IS impressive — say WHAT happened
  and let the reader decide.
- Do not repeat the resume.  Tell the story behind the bullets.

BEFORE AND AFTER EXAMPLES:

Wrong: "I am excited to bring my extensive background in technology delivery
       to your organization."
Right: "I have spent 15 years getting IT programs across the finish line.
       That is what I am good at."

Wrong: "Passionate, results-driven IT leader leveraging 15+ years of progressive
       experience to drive synergistic outcomes."
Right: "I run IT programs from concept through delivery.  Budgets from $250K to
       $3M, across 30+ teams, onshore and offshore."

Wrong: "I look forward to hearing from you at your earliest convenience."
Right: "I would welcome a conversation about what I can bring to [COMPANY]'s
       delivery work."
"""

# ── Job description input ─────────────────────────────────────────────────────

def prompt_for_job_description(company):
    """Interactive paste for CLI use.  Not called when --description is passed."""
    print()
    print("─" * 60)
    print(f"Paste the FULL job description for {company} below.")
    print("The more complete the JD, the better the letter.")
    print("When done, press Enter then type END on its own line.")
    print("─" * 60)
    lines = []
    while True:
        try:
            line = input()
        except EOFError:
            break
        if line.strip().upper() == "END":
            break
        lines.append(line)
    text = "\n".join(lines).strip()
    print(f"Received {len(text)} characters.", flush=True)
    return text or None

# ── Claude API call ──────────────────────────────────────────────────────────

def generate_cover_letter(company, title, job_description, tailored_data=None):
    """Ask Claude to draft the cover letter body.  Returns raw text (no header/sig)."""
    print(f"Generating cover letter for {company} — {title}...", flush=True)

    # Bullet context from tailored JSON
    bullet_context = ""
    if tailored_data:
        bullets    = tailored_data.get("selected_bullets", [])
        key_skills = tailored_data.get("key_skills", [])
        if bullets:
            bullet_context = (
                "\n\nSELECTED RESUME BULLETS FOR THIS ROLE "
                "(reference the facts and numbers in these — do not copy them verbatim):\n"
            )
            for b in bullets:
                bullet_context += f"  [{b.get('company','')}] {b.get('bullet','')}\n"
        if key_skills:
            bullet_context += (
                f"\nKEY SKILLS THAT MATCH THIS ROLE: {', '.join(key_skills)}"
            )

    client = anthropic.Anthropic()

    prompt = f"""You are writing a cover letter FOR Jason W. Darrow.
You must write in his voice and follow every rule below exactly.

{VOICE_RULES}

--- ABOUT JASON ---
{ABOUT_ME}

--- CAREER HISTORY ---
{CAREER_HISTORY}
{bullet_context}

--- THE JOB ---
Company: {company}
Role: {title}

Full job description:
{job_description}

--- WHAT TO WRITE ---
Write the body of Jason's cover letter for this role.  Three to four paragraphs.
No salutation, no signature, no date.  Just the paragraphs.

Paragraph 1 (opening hook):
  Pick ONE of the following approaches — whichever fits best given the JD:
  (a) Open with the specific problem or challenge the role is trying to solve,
      then connect Jason's track record to it.
  (b) Reference something specific and concrete in the JD (a technology, a
      business context, a stated goal) that resonates with Jason's actual work.
  (c) Lead with a single concrete fact from Jason's career that immediately
      answers the question "why does this person belong in this seat."
  Do NOT open with "I", "My", or the company name as the very first word.
  Do NOT open with a compliment about the company.

Paragraph 2 (proof):
  2-3 specific achievements that map directly to what the JD is asking for.
  Use real numbers from the bullets above.  Show causation where it exists.
  Keep sentences short.  Each sentence = one fact or one outcome.

Paragraph 3 (what else Jason brings):
  The AI project angle: Jason built an AI-powered job search system using the
  Claude API and Adzuna API.  He did it to learn AI by doing.  Mention this
  only if it adds genuine value for the role — don't force it.
  His Air Force discipline and delivery mindset if relevant.
  Why he is looking now (Voya layoff 5/4/2026 — his role was eliminated, not
  performance).  Frame it matter-of-factly, one sentence, if at all.

Paragraph 4 (close):
  A specific, confident ask.  Reference {company} by name.
  Do NOT use "I look forward to hearing from you."
  Make it sound like someone who expects the call, not someone who is hoping for it.
"""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}]
    )

    return message.content[0].text.strip()

# ── Word document builder ────────────────────────────────────────────────────

def build_docx(company, title, letter_body, output_path):
    """Write a clean .docx cover letter with header, date, body, signature."""
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    def _para(text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = "Calibri"
        return p

    # ── Name header ─────────────────────────────────────────────────────────
    _para("JASON W. DARROW", bold=True, size=14,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _para(
        "774-573-8354  |  Jason@JasonDarrow.com  |  linkedin.com/in/jasondarrow",
        size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
    )

    # ── Date ────────────────────────────────────────────────────────────────
    _para(date.today().strftime("%B %d, %Y"), size=11, space_after=12)

    # ── Salutation ──────────────────────────────────────────────────────────
    _para(f"Dear Hiring Manager,", size=11, space_after=12)

    # ── Body paragraphs ─────────────────────────────────────────────────────
    # Split on blank lines; each chunk is one paragraph.
    chunks = [c.strip() for c in re.split(r"\n{2,}", letter_body) if c.strip()]
    for chunk in chunks:
        _para(chunk, size=11, space_after=12)

    # ── Signature ───────────────────────────────────────────────────────────
    _para("Sincerely,", size=11, space_after=12)
    _para("Jason W. Darrow", size=11, space_after=0)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"✓ Cover letter saved: {output_path}", flush=True)

# ── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate a tailored cover letter")
    parser.add_argument("--company",     required=True, help="Company name")
    parser.add_argument("--title",       required=True, help="Job title")
    parser.add_argument("--description", help="Full job description text")
    parser.add_argument("--input",       help="Tailored JSON from resume_tailor.py (auto-detected if omitted)")
    parser.add_argument("--output",      help="Output .docx path (default: personal/CoverLetter_COMPANY.docx)")
    args = parser.parse_args()

    # ── Load tailored JSON (auto-detect if not specified) ─────────────────
    tailored_data = None
    input_path = args.input or os.path.join(
        "output", f"tailored_{args.company.replace(' ', '_')}.json"
    )
    if os.path.exists(input_path):
        with open(input_path) as f:
            tailored_data = json.load(f)
        print(f"Loaded tailored data from: {input_path}", flush=True)
    else:
        print(
            f"No tailored JSON found at {input_path} — "
            "proceeding with master resume context only.",
            flush=True,
        )

    # ── Get full job description ───────────────────────────────────────────
    job_description = args.description or None
    if not job_description:
        job_description = prompt_for_job_description(args.company)
    if not job_description or len(job_description) < 100:
        print("Job description too short or missing.  Exiting.")
        exit(1)

    print(f"Job description: {len(job_description)} characters.", flush=True)

    # ── Generate ──────────────────────────────────────────────────────────
    letter_body = generate_cover_letter(
        args.company, args.title, job_description, tailored_data
    )

    # ── Save ──────────────────────────────────────────────────────────────
    output_path = args.output or os.path.join(
        "personal", f"CoverLetter_{args.company.replace(' ', '_')}.docx"
    )
    build_docx(args.company, args.title, letter_body, output_path)

    print("\nDone.", flush=True)
    print(f"File: {output_path}", flush=True)
